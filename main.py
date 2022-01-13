from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.slider import Slider
from kivymd.app import MDApp
from kivy.lang import Builder
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.properties import ObjectProperty
from kivy.metrics import dp
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.stacklayout import StackLayout
from kivy.core.window import Window
from kivymd.toast import toast
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.button import MDRectangleFlatButton
from kivy.uix.button import Button
from kivymd.uix.toolbar import MDToolbar
from kivymd.uix.slider import MDSlider
from kivy.uix.image import Image as Im
from kivy.graphics import Rectangle, Color, Line
from kivymd.uix.label import MDIcon
from kivy.uix.scrollview import ScrollView
from kivy.config import Config
import string
from PIL import Image
from tkinter import filedialog
import os
import numpy as np
import pytesseract
import json
import random
from docx import Document
from threading import Thread
import sys
import tkinter as tk
sys.setrecursionlimit(5000)

app_path = os.path.dirname(os.path.abspath(__file__))
pytesseract.pytesseract.tesseract_cmd=f"{app_path}\\Tesseract-OCR\\tesseract.exe"

Config.set('graphics', 'resizable', True)

sm=None

f=open(f"{app_path}/dict.json",'r')
Dict=json.loads(f.read())
f.close()

f=open(f"{app_path}/setting.json",'r')
Set_dict=json.loads(f.read())
f.close()

f=open(f"{app_path}/image_id.json",'r')
image_id=json.loads(f.read())
f.close()

f=open(f"{app_path}/bgimage.json",'r')
bgimage_id=json.loads(f.read())
f.close()



Image_path=None
BG_Image_path=None
Doc_path=None
Letter_alp=None
Letter_num=None
I=None
M=None
A=None
L=None
S=None
W=None
W_dict={}

bg_image_id=None
Button_id=None
Letter_id=None
table_id=None

BG_but=None
BG_G=None
word_space=int(Set_dict['space'])
letter_gap=int(Set_dict['char'])
default_color=[int(Set_dict['R']),int(Set_dict['G']),int(Set_dict['B'])]
set_bold=int(Set_dict['bold'])

def index():
    for n,i in enumerate(Dict[Letter_alp]):
        if i[0]==Letter_num:
            return n
    return 0

def write(arr,x,y,c,size,color,bold):
    n=random.randint(0,len(Dict[c])-1)
    c_list=Dict[c][n][4]
    specific_size=Dict[c][n][1]
    size*=specific_size


    ya=0
    flag=False
    ja=None
    addy=int(Dict[c][n][3])
    addx=int(Dict[c][n][2])
    
    try:
        i=len(c_list)-1
        while(i>0):
            for j in c_list[int(i)]:
                if flag==False :
                    ja=j
                    flag=True
                else :
                    '''if x+j+1>4500:
                        x=int(right)
                        y+=int(line_gap)'''
                    arr[y-ya-int(size*addy),x+addx+int(ja*size):x+addx+int(j*size)+bold]=color
                    flag=False
                    
            ya+=1
            i-=1/size
        x+=int((Dict[c][n][5]+letter_gap)*size)
    except:
        pass
    return arr,x
    

def Table_generator(path):
    if path.split(".")[-1]=='docx':
        img=Image.open(f"{app_path}/Table/{Set_dict['table']}.jpg")
        img=img.resize((4500,10000))
        file=Document(path)
        table_no=Set_dict['tableno']
        for table in file.tables:
            col=[]
            x_dimmension=4500
            sum_col=0
            for column in table.columns:
                col.append(column.width)
                sum_col+=column.width

            div=(x_dimmension-100)/sum_col
            arr=np.array(img)
            s=50
            col_space=[]
            col_add=[80]
            for i in col:
                y=int(s+i*div)
                col_space.append(y-s)
                col_add.append(y+30)
                s=y

            x=80
            col_add[0]=x

            l_cell=[]
            m_cell=[]
            for row in table.rows:
                l=list(cell for cell in row.cells)
                l_cell.append(l)

            for row in table.columns:
                l=list(cell for cell in row.cells)
                m_cell.append(l)

            mn=[]
            xn=0
            for r in range(len(l_cell)-1,-1,-1):
                ln=[]
                for c in range(len(l_cell[0])-1,-1,-1):

                    if l_cell[r][c]==l_cell[r][c-1]:
                        l_cell[r].pop(c)
                        xn+=1
                    else:
                        ln.insert(0,xn)
                        xn=0

                    if m_cell[c][r]==m_cell[c][r-1]:
                        l_cell[r][c]=False

                ln.insert(0,0)
                mn.insert(0,ln)

            for ni,i in enumerate(mn):
                sum=0
                for nj,j in enumerate(i):
                    sum+=j
                    mn[ni][nj]=sum

    
            max=[0]*(ni+2)
            for m,i1 in enumerate(l_cell): 
                cadd=0
                y=max[m]+200
                y0=y
                x=80
                max[m+1]=y
                for n,cell in enumerate(i1) :
                    try:
                        if l_cell[m+1][n]==False:
                            cadd+=1
                    except:
                        pass
                    if cell==False:
                        continue
                    y=y0
                    for new,para in enumerate(cell.paragraphs) :
                        x=col_add[n+mn[m][n]]
                        if new:
                            y+=170
                            if y>max[m+1+cadd]:
                                max[m+1+cadd]=y
                        word=para.text.split(' ')
                        k=0
                        for run in para.runs :
                            bold=0
                            if run.font.bold :
                                bold=set_bold
                        
                            if run.font.color.rgb !=None:
                                color=[]
                                for i in run.font.color.rgb:
                                    color.append(i)
                            else:
                                color=default_color
                            for c in run.text:
                                if c==' ':
                                    x+=int(word_space*0.8)
                                    k=k+1
                                    if (col_add[n+1+mn[m][n+1]]-x-(len(word[k])*75*Set_dict["size"]))< 0:
                                            x=col_add[n+mn[m][n]]
                                            y+=int(200*Set_dict["size"])
                                            if y>max[m+1+cadd]:
                                                max[m+1+cadd]=y
                                    continue
                                else:
                                    try:
                                        arr,x=write(arr,x,y,c,Set_dict["size"],color,bold)
                                    except:
                                        pass
                                
                            
            for m,i in enumerate(max) :
                for n in range(len(col_add)-1):
                    try:
                        if l_cell[m][n]==False:
                            continue
                    except:
                        pass
                    arr[i+50:i+60,col_add[n]-30:col_add[n+1]-30]=[83,83,83] 

            for n,j in enumerate(l_cell):
                for m in range(len(l_cell[n])+1):
                    try:
                        arr[max[n]+50:max[n+1]+50,col_add[m+mn[n][m]]-30:n+col_add[m+mn[n][m]]-20]=[83,83,83]       
                    except:
                        pass

            arr=arr[:max[-1]+110,:]
            pageimg = Image.fromarray(arr)
            pageimg=pageimg.resize((pageimg.width*Set_dict['img_size']//100,pageimg.height*Set_dict['img_size']//100))
            pageimg.save(f"{app_path}/work/table-{table_no}.jpg")
            table_no+=1
        Set_dict['tableno']=table_no
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        sm.current='Main'
        toast('Work Done!! Check work folder')
    else :
        toast('Not a valid docx file')
    

def Passage_generator(path): 
    if path.split(".")[-1]=='docx':

        for n,i in enumerate(bgimage_id):
            if Set_dict['bg']==i[0]:
                index=n
                break
        
        
        img=Image.open(f"{app_path}/BGimage/{Set_dict['bg']}.jpg")
        arr=np.array(img)
        file=Document(path)
        lines=file.paragraphs
        
        top=int(bgimage_id[index][1])
        bottom=int(bgimage_id[index][2])
        right=int(bgimage_id[index][3])
        left=int(bgimage_id[index][4])
        line_gap=int(bgimage_id[index][5])
        

        page_no=Set_dict['pageno']
        underline=False
        y=int(top+line_gap)
        superscript=False
        color=default_color

        font_dic={}
        for line in lines:
            for diff in line.runs:
                try:
                    font=diff.font.size.pt
                except :
                    pass
                if font_dic.get(font,False)==False :
                    font_dic[font]=len(diff.text)
                else :
                    font_dic[font]+=len(diff.text)
        val_list=list(font_dic.values())
        key_list=list(font_dic.keys())
        position = val_list.index(max(val_list))

        basic_font_size=key_list[position]

        for line in lines:
            if underline :
                arr[y+10:y+16,underline_x:x]=color
            x=right
            y+=line_gap
            k=0
            
    
            length=line.text.split(' ') 
            underline=False
            for diff in line.runs :  
                bold=0
                try:
                    Font_size=diff.font.size.pt
                    font_size=Font_size/basic_font_size
                except: 
                    font_size=1
                
                if diff.font.bold :
                    bold=set_bold

                if diff.font.color.rgb !=None:
                    color=[]
                    for i in diff.font.color.rgb:
                        color.append(i)
                else:
                    color=default_color

                if underline :
                    arr[y+10:y+16,underline_x:x]=color
                    underline=False
                if diff.underline!=None:
                    underline=True
                    underline_x=x
                if superscript:
                    superscript=False
                    y+=int(80*Font_size/basic_font_size)
                if diff.font.superscript:
                    y-=int(80*Font_size//basic_font_size)
                    Font_size=Font_size//1.5
                    superscript=True

                if diff.font.subscript:
                    Font_size=Font_size//1.5


                
                for c in diff.text:
                    if c==" " :
                        k+=1
                        x+=int(word_space)
                        if ((4500-x-left)-(len(length[k])*75*Set_dict["size"]*font_size)) < 0:
                            if underline:
                                arr[y+10:y+16,right:x-int(word_space)]=color
                            y+=int(line_gap)
                            x=int(right)
                            if y>(img.height-bottom):
                                y=int(top+line_gap)
                                imag=Image.fromarray(arr)
                                imag=imag.resize((imag.width*Set_dict['img_size']//100,imag.height*Set_dict['img_size']//100))
                                imag.save(f"{app_path}/work/page-{page_no}.jpg")
                                page_no+=1
                                del arr
                                arr=np.array(img)

                    elif Dict.get(c,False) and len(Dict[c])>0:
                        arr,x=write(arr,x,y,c,(Set_dict["size"]*font_size),color,bold)
                        
        imag=Image.fromarray(arr)
        imag=imag.resize((imag.width*Set_dict['img_size']//100,imag.height*Set_dict['img_size']//100))
        imag.save(f"{app_path}/work/page-{page_no}.jpg")
        Set_dict['pageno']=page_no+1
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        sm.current='Main'
        toast('Work Done!! Check work folder')
    else :
        toast('Not a valid docx file')
    


def add_l(path):
    ext = path.split(".")[-1]
    if ext == "jpeg" or ext == "jpg" or ext == "png":
        img = Image.open(path)
        name=pytesseract.image_to_boxes(img)
        width,height=img.size
        img_arr=np.array(img)
        num=Set_dict['num']
        for line in name.splitlines():
            x=line.split(' ')
            if x[0]=='~' or x[0]=='@':
                continue
            l,t,w,h=int(x[1]),height-int(x[2]),int(x[3]),height-int(x[4])
            c=[]
            for i in range(h,t):
                m=[]
                flag=False
                for j in range(l,w):
                    if img_arr[i,j,0]<100 and img_arr[i,j,1]<150 :
                        if flag==False:
                            m.append(j-l)
                            flag=True
                    else:
                        img_arr[i,j]=[255,255,255]
                        if flag==True :
                            m.append(j-l)
                            flag=False
                if flag:
                    m.append(j-l)    
                c.append(m)
            if len(c)==0:
                continue
            try:
                num+=1
                text=x[0]
                yl=0
                list_text='fgjpqy'
                for i in list_text:
                    if i == text:
                        yl=-(t-h)//2
                        break
                if Dict.get(text,False)==False:
                    Dict[text]=[]
                Dict[text].append([num,1,0,yl,c,w-l])
                a=img_arr[h:t,l:w]
                imga=Image.fromarray(a)
                i=text
                imga.save(f"{app_path}/images/{num}.jpg")
                image_id[str(num)]=text
            except:
                pass
        Set_dict['num']=num
        with open(f"{app_path}/dict.json", "w") as f:
            json.dump(Dict, f)
        f.close()
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        with open(f"{app_path}/image_id.json", "w") as f:
            json.dump(image_id, f)
        f.close()
        global S,W
        S.remove_widget(W)
        S.add_widget(Write())
    else :
        toast('Not a valid image file')

def filter_image(path,text=None):
    ext = path.split(".")[-1]
    if ext == "jpeg" or ext == "jpg" or ext == "png":
        img = Image.open(path)
        width, height = img.size
        height=75*height//width
        img=img.resize((75,height))
        img_arr=np.array(img)
        c=[]
        for i in range(height):
            m=[]
            flag=False
            for j in range(75):
                if img_arr[i,j,0]<100 and img_arr[i,j,1]<150:
                    if flag==False:
                        m.append(j)
                        flag=True
                    elif j==74 :
                        m.append(j)
                        break
                else:
                    img_arr[i,j]=[255,255,255]
                    if flag==True :
                        m.append(j)
                        flag=False
            c.append(m)
        Set_dict['num']+=1
        list_text='fgjpqy'
        yl=0
        for i in list_text:
            if i == text:
                yl=-height//2
                break
        Dict[text].append([Set_dict['num'],1,0,yl,c,75])
        img=Image.fromarray(img_arr)
        img.save(f"{app_path}/images/{Set_dict['num']}.jpg")
        image_id[str(Set_dict['num'])]=text
        W_dict[text].add_widget(Letter_button(text,len(Dict[text])-1, W_dict[text]))
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        with open(f"{app_path}/image_id.json", "w") as f:
            json.dump(image_id, f)
        f.close()
        with open(f"{app_path}/dict.json", "w") as f:
            json.dump(Dict, f)
        f.close()
    else :
        toast('Not a valid image file')
        
class toolbar(MDToolbar,Thread):
    pass

class Itoolbar(MDToolbar,Thread):
    def __init__(self,path=None,**kwargs):
        super().__init__(**kwargs)
        self.left_action_items=[['arrow-left-thick',self.back]]
        self.title= 'Preview'
        self.md_bg_color=( 0.125, 0.125, 0.125, 1)
        self.specific_text_color=( 1, 1, 1, 1)

    def back(self,*args):
        sm.current='BGimage'

class Image_Select(Screen,Thread):
    def __init__(self,path=None,**kwargs):
        super().__init__(**kwargs)
        self.path=path
        self.num=int(path.split("/")[-1].split(".")[-2])

        for n,a in enumerate(bgimage_id):
            if a[0]==self.num:
                self.ind=n
                break
        s=BoxLayout(orientation='vertical')
        b=BoxLayout(spacing=20)
        b.height=Window.width*2.8
        s.add_widget(Itoolbar())
        I=Im(source=path)
        b.add_widget(I)

        d=BoxLayout(orientation='vertical',size_hint=(1,1))
        d.add_widget(Label(text='Alignment'))

        e=BoxLayout(spacing=-4)
        e.add_widget(Label(text='Top'))
        e.add_widget(MDIcon(icon='format-vertical-align-top'))
        self.es=Slider(min=0,max=1500,size_hint=(4,1),value=bgimage_id[self.ind][1])
        self.es.bind(value=self.on_value_e)
        e.add_widget(self.es)
        self.el=Label(text=str(int(bgimage_id[self.ind][1])))
        e.add_widget(self.el)
        d.add_widget(e)

        f=BoxLayout(spacing=-4)
        f.add_widget(Label(text='Bottom'))
        f.add_widget(MDIcon(icon='format-vertical-align-bottom'))
        self.fs=Slider(min=0,max=1500,size_hint=(4,1),value=bgimage_id[self.ind][2])
        self.fs.bind(value=self.on_value_f)
        f.add_widget(self.fs)
        self.fl=Label(text=str(int(bgimage_id[self.ind][2])))
        f.add_widget(self.fl)
        d.add_widget(f)

        g=BoxLayout(spacing=-4)
        g.add_widget(Label(text='Right'))
        g.add_widget(MDIcon(icon='format-horizontal-align-right'))
        self.gs=Slider(min=0,max=1000,size_hint=(4,1),value=bgimage_id[self.ind][3])
        self.gs.bind(value=self.on_value_g)
        g.add_widget(self.gs)
        self.gl=Label(text=str(int(bgimage_id[self.ind][3])))
        g.add_widget(self.gl)
        d.add_widget(g)

        h=BoxLayout(spacing=-4)
        h.add_widget(Label(text='Left'))
        h.add_widget(MDIcon(icon='format-horizontal-align-left'))
        self.hs=Slider(min=0,max=1000,size_hint=(4,1),value=bgimage_id[self.ind][4])
        self.hs.bind(value=self.on_value_h)
        h.add_widget(self.hs)
        self.hl=Label(text=str(int(bgimage_id[self.ind][4])))
        h.add_widget(self.hl)
        d.add_widget(h)


        i=BoxLayout(orientation='vertical',padding=(0,20,0,20),size_hint=(1,2))
        i.add_widget(Label(text='Distance betwwen Line'))
        j=BoxLayout()
        j.add_widget(MDIcon(icon='arrow-expand-vertical'))
        self.js=Slider(min=0,max=500,size_hint=(4,1),value=bgimage_id[self.ind][5])
        self.js.bind(value=self.on_value_j)
        j.add_widget(self.js)
        self.jl=Label(text=str(int(bgimage_id[self.ind][5])))
        j.add_widget(self.jl)
        i.add_widget(j)

        d.add_widget(i)

        k=BoxLayout()
        de=Button(text='Delete',on_release=self.delete)
        k.add_widget(de)
        k.add_widget(Button(text='Cancel',on_release=self.cancel))
        k.add_widget(Button(text='Save',on_release=self.save))
        k.add_widget(Button(text='Select',on_release=self.select))
        d.add_widget(k)
        b.add_widget(d)
        s.add_widget(b)
        
        self.add_widget(s)

    def on_value_e(self,instance,val):
        self.el.text="%d" %val

    def on_value_f(self,instance,val):
        self.fl.text="%d" %val
    
    def on_value_g(self,instance,val):
        self.gl.text="%d" %val
    
    def on_value_h(self,instance,val):
        self.hl.text="%d" %val

    def on_value_j(self,instance,val):
        self.jl.text="%d" %val

    def cancel(self,*args):
        sm.current='BGimage'

    def delete(self,*args):
        if len(bgimage_id)==1:
            toast("There should be atleast 1 image in passage")
        else:    
            bgimage_id.pop(self.ind)
            if self.num==Set_dict['bg']:
                Set_dict['bg']=bgimage_id[0][0]
                bg_image_id.background_normal=f"{app_path}/BGimage/"+str(bgimage_id[0][0])+".jpg"
                with open(f"{app_path}/setting.json", "w") as f:
                    json.dump(Set_dict, f)
                f.close()
            with open(f"{app_path}/bgimage.json", "w") as f:
                json.dump(bgimage_id, f)
            f.close()
            os.remove(self.path)
            BG_G.remove_widget(BG_but)
            sm.current='BGimage'
    
    def save(self,*args):
        bgimage_id[self.ind][1]=self.es.value
        bgimage_id[self.ind][2]=self.fs.value
        bgimage_id[self.ind][3]=self.gs.value
        bgimage_id[self.ind][4]=self.hs.value
        bgimage_id[self.ind][5]=self.js.value

        with open(f"{app_path}/bgimage.json", "w") as f:
            json.dump(bgimage_id, f)
        f.close()

        sm.current='BGimage'

    def select(self,*args):
        Set_dict['bg']=self.num
        bg_image_id.background_normal=self.path
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        self.save(*args)

class Btoolbar(MDToolbar,Thread):
    pass

   
class BGimage(Screen,Thread):


    def __init__(self,**kwargs):
        super().__init__(**kwargs)
        box=BoxLayout(orientation='vertical')
        B=Btoolbar()
        box.add_widget(B)
        l=len(bgimage_id)
        S=ScrollView()
        G=StackLayout(padding=10,spacing=5)
        global BG_G
        BG_G=G
        
        num=1
        while (l>=num):
            self.b=Button(color=(0,0,0,1),background_normal=f"{app_path}/BGimage/{bgimage_id[num-1][0]}.jpg",border=(0,0,0,0),on_release=self.change,
            size_hint=(None,None),size=(150,200))
            num+=1
            G.add_widget(self.b)
        num-=1
        S.add_widget(G)
        box.add_widget(S)
        self.add_widget(box)


    def change(self,instance):
        global I
        try:
            sm.remove_widget(I)
        except:
            pass
        I=Image_Select(path=instance.background_normal,name='1')
        sm.add_widget(I)
        sm.current='1'
        global BG_Image_path,BG_but
        BG_but=instance
        BG_Image_path=instance.background_normal

    def file_manager_open(self,*args):
        root = tk.Tk()
        root.withdraw()

        files = filedialog.askopenfilenames()
        for i in files:
            self.select_path(i)


    def select_path(self, path):
        ext=path.split(".")[-1]
        if ext=='jpg' or ext=='jpeg' or ext=='png':
            global BG_Image_path
            BG_Image_path=path
            img = Image.open(path)
            img=img.resize((4500,int(4500*(img.height/img.width))))
            img.save(f"{app_path}/BGimage/{Set_dict['bg_max']}.jpg")
            bgimage_id.append([Set_dict['bg_max'],200,200,200,200,170])
            Set_dict['bg_max']+=1
            with open(f"{app_path}/setting.json", "w") as f:
                json.dump(Set_dict, f)
            f.close()
            with open(f"{app_path}/bgimage.json", "w") as f:
                json.dump(bgimage_id, f)
            f.close()
            BG_G.add_widget(Button(color=(0,0,0,1),background_normal=f"{app_path}/BGimage/{Set_dict['bg_max']-1}.jpg",border=(0,0,0,0),on_release=self.change,
            size_hint=(None,None),size=((Window.width-30)/3,(Window.width-30)*(4/9))))
        else:
            toast("Try Again With Valid Image")


Table_img=""
G_table=None
img_path=None

class TableImagePreview(Screen,Thread):  
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        B=BoxLayout(orientation='vertical')
        
        tool=MDToolbar(title='Preview',md_bg_color=( 0.125, 0.125, 0.125, 1),specific_text_color=( 1, 1, 1, 1))
        tool.left_action_items=[['arrow-left-thick',self.change_screen]]
        B.add_widget(tool)
        global Table_img,img_path
        img_path=Im()
        B.add_widget(img_path)


        Box=BoxLayout(size_hint=(1,None),height=Window.height*0.1)
        Box.add_widget(Button(text='Delete',on_release=self.delete))
        Box.add_widget(Button(text='Cancel',on_release=self.change_screen))
        Box.add_widget(Button(text='Select',on_release=self.select))
        B.add_widget(Box)
        self.add_widget(B)

    def change_screen(self,*args):
        sm.current='TableImage'
    
    def delete(self,*args):

        arr=os.listdir(f"{app_path}/Table/")
        if len(arr)==1:
            toast("There should be atleast 1 image in table")
        else:  
            try:
                ids=int(img_path.source.split("/")[-1].split('.')[-2])
                sm.current='TableImage'
                G_table.remove_widget(Table_img)
                os.remove(img_path.source)
                if ids==Set_dict['table']:
                    Set_dict['table']=int(os.listdir(f"{app_path}/Table/")[0].split('.')[-2])
                    table_id.background_normal=f"{app_path}/Table/{Set_dict['table']}.jpg"
                with open(f"{app_path}/setting.json", "w") as f:
                    json.dump(Set_dict, f)
                f.close()
            except:
                pass
    
    def select(self,*args):
        sm.current='TableImage'
        Set_dict['table']=int(img_path.source.split("/")[-1].split('.')[-2])
        with open(f"{app_path}/setting.json", "w") as f:
            json.dump(Set_dict, f)
        f.close()
        table_id.background_normal=img_path.source

class TableImage(Screen,Thread):
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        B=BoxLayout(orientation='vertical')
        
        tool=MDToolbar(title='Table',md_bg_color=( 0.125, 0.125, 0.125, 1),specific_text_color=( 1, 1, 1, 1))
        tool.left_action_items=[['arrow-left-thick',self.change_screen]]
        S=ScrollView()
        B.add_widget(tool)
        global G_table
        G_table=StackLayout(padding=10,spacing=5)
        arr=os.listdir(f"{app_path}/Table/")
        for i in arr: 
            self.b=Button(color=(0,0,0,1),background_normal=f"{app_path}/Table/{i}",border=(0,0,0,0),on_release=self.change,
            size_hint=(None,None),size=(150,200))
            G_table.add_widget(self.b)
        S.add_widget(G_table)
        B.add_widget(S)

        self.add_widget(B)

    def change_screen(self,*args):
        sm.current='Main'

    def change(self,instance):
        global Table_img,img_path,boo
        Table_img=instance
        img_path.source=Table_img.background_normal
        sm.current='TablePreview'
        

    def file_manager_open(self,*args):
        root = tk.Tk()
        root.withdraw()

        files = filedialog.askopenfilenames()
        for i in files:
            self.select_path(i)


    def select_path(self, path):
        ext=path.split(".")[-1]
        if ext=='jpg' or ext=='jpeg' or ext=='png':
            img = Image.open(path)
            img.save(f"{app_path}/Table/{Set_dict['tb_max']}.jpg")
            Set_dict['tb_max']+=1
            with open(f"{app_path}/setting.json", "w") as f:
                json.dump(Set_dict, f)
            f.close()
            G_table.add_widget(Button(color=(0,0,0,1),background_normal=f"{app_path}/Table/{Set_dict['tb_max']-1}.jpg",border=(0,0,0,0),on_release=self.change,
            size_hint=(None,None),size=((Window.width-30)/3,(Window.width-30)*(4/9))))
        else:
            toast("Try Again With Valid Image File")

class Setting(BoxLayout,Thread):
    def __init__(self,**kwargs):
        super().__init__(**kwargs)
        self.size_hint=(1,1)
        self.change=True
        b=BoxLayout(orientation='vertical')
        b.size_hint=(0.5,1)
        b.spacing=5
        b.padding=10
        global bg_image_id,table_id
        bg_image_id=Button(text='Passage',size_hint=(0.7,1),pos_hint={'center_x':0.5},color=(0,0,0,1),background_normal=f"{app_path}/BGimage/{Set_dict['bg']}.jpg",on_release=self.bg_image,border=(0,0,0,0))
        b.add_widget(bg_image_id)
        table_id=Button(text='Table',size_hint=(0.7,1),pos_hint={'center_x':0.5},color=(0,0,0,1),background_normal=f"{app_path}/Table/{Set_dict['table']}.jpg",border=(0,0,0,0),on_release=self.tableimage)
        b.add_widget(table_id)
        
        self.add_widget(b)
        ver=BoxLayout(orientation='vertical')
        
        s=BoxLayout()
        fs=Label(text='Size',font_size="17sp",size_hint=(0.25,1),width=50)
        s.add_widget(fs)
        self.s_sli=Slider(min=0,max=2,step=0.1,size_hint=(1,1),value=Set_dict["size"])
        self.s_sli.bind(value=self.on_value_s)
        self.s_lab=Label(text=str(Set_dict["size"])[:3],size_hint=(0.2,1))
        s.add_widget(self.s_sli)
        s.add_widget(self.s_lab)
        ver.add_widget(s)

        h=BoxLayout()
        h.add_widget(Label(text='Space',font_size="17sp",size_hint=(0.25,1),width=50))
        self.h_sli=Slider(min=0,max=200,step=1,size_hint=(1,1),value=Set_dict["space"])
        self.h_sli.bind(value=self.on_value_h)
        self.h_lab=Label(text=str(int(Set_dict["space"])),size_hint=(0.2,1))
        h.add_widget(self.h_sli)
        h.add_widget(self.h_lab)
        ver.add_widget(h)

        v=BoxLayout()
        v.add_widget(Label(text='Letter Gap',font_size="17sp",size_hint=(0.25,1)))
        self.v_sli=Slider(min=0,max=50,step=1,size_hint=(1,1),value=Set_dict["char"])
        self.v_sli.bind(value=self.on_value_v)
        self.v_lab=Label(text=str(Set_dict["char"]),size_hint=(0.2,1))
        v.add_widget(self.v_sli)
        v.add_widget(self.v_lab)
        ver.add_widget(v)

        
        bo=BoxLayout()
        bo.add_widget(Label(text='Bold',font_size="17sp",size_hint=(0.25,1),width=50))
        self.bo_sli=Slider(min=0,max=20,step=1,size_hint=(1,1),value=Set_dict["bold"])
        self.bo_sli.bind(value=self.on_value_bo)
        self.bo_lab=Label(text=str(Set_dict["bold"]),size_hint=(0.2,1))
        bo.add_widget(self.bo_sli)
        bo.add_widget(self.bo_lab)
        ver.add_widget(bo)

        lo=BoxLayout()
        lo.add_widget(Label(text='Image Size',font_size="17sp",size_hint=(0.25,1),width=50))
        self.lo_sli=Slider(min=1,max=100,step=1,size_hint=(1,1),value=Set_dict["img_size"])
        self.lo_sli.bind(value=self.on_value_lo)
        self.lo_lab=Label(text=str(int(Set_dict["img_size"]))+'%',size_hint=(0.2,1))
        lo.add_widget(self.lo_sli)
        lo.add_widget(self.lo_lab)
        ver.add_widget(lo)

        co=BoxLayout(orientation='vertical',size_hint=(1,2))
        co.add_widget(Label(text='Default Font Color',font_size="17sp",size_hint=(0.25,1),pos_hint={'center_x':0.5},width=50))
        self.R=MDSlider(min=0,max=255,step=1,color=(1,0,0,1),hint_text_color=(1,1,1,1),value=Set_dict["R"])
        self.G=MDSlider(min=0,max=255,step=1,color=(0,1,0,1),hint_text_color=(1,1,1,1),value=Set_dict["G"])
        self.B=MDSlider(min=0,max=255,step=1,color=(0,0,1,1),hint_text_color=(1,1,1,1),value=Set_dict["B"])
        self.R.bind(value=self.on_value_R)
        self.G.bind(value=self.on_value_G)
        self.B.bind(value=self.on_value_B)
        co.add_widget(self.R)
        co.add_widget(self.G)
        co.add_widget(self.B)
        ver.add_widget(co)

        self.do=BoxLayout(size_hint=(1,None),height=65,padding=10)
        self.save_b=Button(text='Save',on_release=self.save,disabled=True)
        self.do.add_widget(self.save_b)
        self.cancel_b=Button(text='Cancel',on_release=self.cancel,disabled=True)
        self.do.add_widget(self.cancel_b)
        ver.add_widget(self.do)
        self.add_widget(ver)
        
    def on_value_R(self, instance, val):
        if self.change:
            self.changes()
    def on_value_G(self, instance, val):
        if self.change:
            self.changes()
    def on_value_B(self, instance, val):
        if self.change:
            self.changes()

    def on_value_s(self, instance, val):
        self.s_lab.text = str(val)[:3]
        if self.change:
            self.changes()
    
    def on_value_v(self, instance, val):
        self.v_lab.text = str(val)[:3]
        if self.change:
            self.changes()
            
    
    def on_value_h(self, instance, val):
        self.h_lab.text = "% d"% val
        if self.change:
            self.changes()

    def on_value_lo(self, instance, val):
        self.lo_lab.text = str(int(val))+'%'
        if self.change:
            self.changes()

    def on_value_bo(self, instance, val):
        self.bo_lab.text = "% d"% val
        if self.change:
            self.changes()

    def changes(self):
        self.change=False
        self.save_b.disabled=False
        self.cancel_b.disabled=False

    def cancel(self,*args):
        self.s_sli.value=Set_dict["size"]
        self.h_sli.value=Set_dict["space"]
        self.v_sli.value=Set_dict["char"]
        self.bo_sli.value=Set_dict["bold"]
        self.lo_sli.value=Set_dict["img_size"]
        self.R.value=Set_dict["R"]
        self.G.value=Set_dict["G"]
        self.B.value=Set_dict["B"]
        self.change=True
        self.save_b.disabled=True
        self.cancel_b.disabled=True

    def save(self,*args):
        Set_dict["size"]=self.s_sli.value
        Set_dict["space"]=self.h_sli.value
        Set_dict["char"]=self.v_sli.value
        Set_dict["bold"]= int(self.bo_sli.value)
        Set_dict["img_size"]= int(self.lo_sli.value)
        Set_dict["R"]= self.R.value
        Set_dict["G"]= self.G.value
        Set_dict["B"]= self.B.value
        self.save_b.disabled=True
        self.cancel_b.disabled=True
        self.change=True
        with open(f"{app_path}/setting.json", "w") as f:
                json.dump(Set_dict, f)
        f.close()
        toast("Saved Sucessfuly")

    def bg_image(self,*args):
        global sm
        sm.current='BGimage'
    
    def tableimage(self,*args):
        sm.current='TableImage'

class Manager(ScreenManager,Thread):
    Menu = ObjectProperty(None)
    About = ObjectProperty(None)

    def change_screen(self):
        self.current='Menu'
    def change_screen_About(self):
        self.current='Menu'

class Image_Box(BoxLayout,Thread):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.index=index()
        self.sizey=0.15
        self.orientation='vertical'
        self.add_widget(toolbar())
        img = Image.open(Image_path)
        wi,he=img.size
        self.hl=he*0.2*Window.height/170
        self.wl=self.hl*wi/he
        img.close()
        self.I=Im(source=Image_path,allow_stretch = True,keep_ratio =True,size_hint=(None,None),
                  height=self.hl*Dict[Letter_alp][self.index][1]*Set_dict['size'],width=self.wl*Dict[Letter_alp][self.index][1]*Set_dict['size'],
                  pos_hint={'y':1/3+Dict[Letter_alp][self.index][3]*(0.2/0.45)/170,'center_x':0.5})
        self.B=BoxLayout(size_hint=(1,None),height=0.45*Window.height,padding=((Window.width-self.I.width)//2,0,0,0))
        self.B.add_widget(self.I)
        self.add_widget(self.B)

        sl=BoxLayout(orientation='vertical')
        t=BoxLayout(padding=10,size_hint=(None,1),width=100,pos_hint={'center_x':0.5})
        t.add_widget(MDIcon(icon='alpha-t-box-outline',size_hint=(None,1),width=40))
        self.T=TextInput(text=Letter_alp,foreground_color = (1,1,1,1),size_hint=(1,1), multiline=False,background_color=(0.2,0.2,0.2,1))
        t.add_widget(self.T)
        sl.add_widget(t)
     
        s=BoxLayout()
        s.add_widget(MDIcon(icon='resize',size_hint=(None,1),width=40))
        self.s_sli=Slider(min=0,max=2,step=0.1,size_hint=(1,1),value=Dict[Letter_alp][self.index][1])
        self.s_sli.bind(value=self.on_value_s)
        self.s_lab=Label(text=str(Dict[Letter_alp][self.index][1])[:3],size_hint=(0.2,1))
        s.add_widget(self.s_sli)
        s.add_widget(self.s_lab)
        sl.add_widget(s)


        h=BoxLayout()
        h.add_widget(MDIcon(icon='format-horizontal-align-center',size_hint=(None,1),width=40))
        self.h_sli=Slider(min=-80,max=80,step=1,size_hint=(1,1),value=Dict[Letter_alp][self.index][2])
        self.h_sli.bind(value=self.on_value_h)
        self.h_lab=Label(text=str(Dict[Letter_alp][self.index][2]),size_hint=(0.2,1))
        h.add_widget(self.h_sli)
        h.add_widget(self.h_lab)
        sl.add_widget(h)

        v=BoxLayout()
        v.add_widget(MDIcon(icon='format-vertical-align-center',size_hint=(None,1),width=40,pos_hint={'center_x':0.5}))
        self.v_sli=Slider(min=-110,max=110,step=1,size_hint=(1,1),value=Dict[Letter_alp][self.index][3])
        self.v_sli.bind(value=self.on_value_v)
        self.v_lab=Label(text=str(Dict[Letter_alp][self.index][3]),size_hint=(0.2,1))
        v.add_widget(self.v_sli)
        v.add_widget(self.v_lab)
        sl.add_widget(v)


        b=MDBoxLayout()
        b_del=Button(text='Delete')
    
        if Image_path and os.path.isfile(Image_path):
            b_del.bind(on_release=self.delete)
        b.add_widget(b_del)
        b.add_widget(Button(text='Cancel',on_release=self.cancel))
        b.add_widget(Button(text='Save',on_release=self.save))
        
        sl.add_widget(b)
        self.add_widget(sl)

    def on_value_s(self, instance, val):
        self.s_lab.text = str(val)[:3]
        self.I.height=self.hl*val
        self.I.width=self.wl*val
    
    def on_value_h(self, instance, val):
        self.h_lab.text = "% d"% val
        self.B.padding=((Window.width-self.I.width)//2+val*(0.2*Window.height)/170,0,0,0)

    def on_value_v(self, instance, val):
        self.v_lab.text = "% d"% val
        self.I.pos_hint={'y':1/3+val*(0.2/0.45)/170}

    def cancel(self,*args):
        sm.current='Main'


    def save(self,*args):

        Dict[Letter_alp][self.index][1]=self.s_sli.value
        Dict[Letter_alp][self.index][2]=self.h_sli.value
        Dict[Letter_alp][self.index][3]=self.v_sli.value

        if self.T.text !=Letter_alp:
            if Dict.get(self.T.text,False)==False:
                toast("'"+self.T.text+"' character not found!!")
            else:
                Dict[self.T.text].append(Dict[Letter_alp][self.index])
                Dict[Letter_alp].pop(self.index)
                image_id.pop(str(Letter_num))
                image_id[str(Letter_num)]=self.T.text
                with open(f"{app_path}/image_id.json", "w") as f:
                    json.dump(image_id, f)
                A.remove_widget(Button_id)
                W_dict[self.T.text].add_widget(Letter_button(self.T.text,len(Dict[self.T.text])-1, W_dict[self.T.text]))
                f.close()

        with open(f"{app_path}/dict.json", "w") as f:
            json.dump(Dict, f)
        f.close()
        sm.current='Main'

    def delete(self,*args):
        Dict[Letter_alp].pop(self.index)
        with open(f"{app_path}/dict.json", "w") as f:
            json.dump(Dict, f)
        f.close()
        global Button_id
        A.remove_widget(Button_id)
        os.remove(Image_path)
        sm.current='Main'

class Image_Choose(Screen,Thread):
    pass

class Image_Screen(Screen,Thread):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.add_widget(Image_Box())
        with self.canvas.before:
            Rectangle(pos=(0,Window.height*0.55-64),size=(Window.width,Window.height*0.45))
        with self.canvas:
            Color(0,0,1)
            Line(points=(0,Window.height*0.7-64,Window.width,Window.height*0.7-64),width=0.2)
            Line(points=(0,Window.height*0.9-64,Window.width,Window.height*0.9-64),width=0.2)
        

    def change_screen(self):
        MDApp.get_running_app().root.current='Main'

class Letter_button(Button,Thread):

    def __init__(self,i,f,alpha, **kwargs):
        super().__init__(**kwargs)
        self.alpha=alpha
        self.size_hint=(None,None)
        self.size=("45dp","45dp")
        self.pos_hint={'y': 0, 'x': 0}
        self.background_normal = f"{app_path}/images/{Dict[i][f][0]}.jpg"
        self.border=(0, 0, 0, 0)
        self.on_release=self.change_screen
        
    
    def change_screen(self,*args):

        global I,A,Image_path,Letter_alp,Letter_num,Button_id

        Image_path=str(self.background_normal)
        Letter_alp=image_id[Image_path.split("/")[-1].split(".")[-2]]
        Letter_num=int(Image_path.split("/")[-1].split(".")[-2])
        Button_id=self
        A=self.alpha
        try:
            sm.remove_widget(I)
        except:
            pass
        I=Image_Screen(name='p')
        sm.add_widget(I)
        sm.current='p'
   
        
class Letter(BoxLayout,Thread):

    manager_open = False
    file_manager = None


    def __init__(self, i, **kwargs):
        super().__init__(**kwargs)
        self.size_hint = (1, None)
        global L,W_dict
        L=self
        self.size_hint = (1, None)
        self.height=50

        self.text=i
        self.k=1
        self.padding=10
        self.spacing=8
        b1 = MDRectangleFlatButton(text=i, font_size="20sp", theme_text_color="Custom", md_bg_color=(
            0, 0, 1, 1), text_color=(1, 1, 1, 1), pos_hint={'top': 1, 'x': 0}, on_release = self.file_manager_open,
            size_hint_y=None,height="45dp")
        self.add_widget(b1)
        
        self.alpha=BoxLayout(spacing=5,padding=(0,-15,0,-15))
        W_dict[i]=self.alpha
        self.num=1
        try:
            l=len(Dict[i])
            f=0
            while(f<l):
                self.b2 = Letter_button(i,f,self.alpha)         
                self.alpha.add_widget(self.b2)
                f+=1
        except:
            pass
        self.add_widget(self.alpha)




    def file_manager_open(self,*args):
        root = tk.Tk()
        root.withdraw()

        files = filedialog.askopenfilenames()
        for i in files:
            filter_image(i,self.text)
        
class Write(BoxLayout,Thread):
     def __init__(self, **kwargs):
        super().__init__(**kwargs) 
        global W
        W=self
        self.orientation = 'vertical'
        self.size_hint = (1, None)
        s = string.ascii_lowercase+string.ascii_uppercase
        for i in range(10):
            s += str(i)
        s+=r".,:;'+-*=^<>%()"
        s+=u"\u00F7"
        s+=r"{}#@&$!"

        for i in s:
            try:
                if Dict.get(i,False)==False:
                    Dict[i]=[]
                b1 = Letter(i)
                self.add_widget(b1)
            except:
                pass
        self.height=50*len(s)


class Writing(ScrollView,Thread):
    def __init__(self, **kwargs):
        super().__init__(**kwargs) 
        global S
        S=self
        self.add_widget(Write())
        

class MenuScreen(Screen,Thread):
    pass

class About_Screen(Screen,Thread):
    pass


class MainScreen(Screen,Thread):
    def __init__(self,**kwargs):
        super().__init__(**kwargs)
        global M
        M=self



    def py_test(self):
        root = tk.Tk()
        root.withdraw()

        file = filedialog.askopenfilename()
        add_l(file)

    def path(self):
        return Image_path



class Scr(ScreenManager,Thread):

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        global sm 
        sm=self

    def change_screen(self):
        self.current='Main'

        
    def doc_file_manager(self,*args):
        root = tk.Tk()
        root.withdraw()

        file = filedialog.askopenfilename()
        Passage_generator(file)

    def file_manager_open(self,*args):
        root = tk.Tk()
        root.withdraw()

        file = filedialog.askopenfilename()
        Table_generator(file)


    def path(self):
        return Image_path

class Front_BoxLayout(BoxLayout,Thread):
    def __init__(self,**kwargs):
        super().__init__(**kwargs)
        self.padding="6dp"
        self.spacing="3dp"
        b=Button(text="PASSAGE",font_size="40sp",color=(0,0,1,1),border=(0, 0, 0, 0),background_normal = f"{app_path}/front/pass2.jpg",
        on_press=self.change)   
        self.add_widget(b)
        b1=Button(text="TABLE",font_size="40sp",color=(0.75,0,0,1),border=(0, 0, 0, 0),background_normal = f"{app_path}/front/table.jpg",
        on_press=self.changet)
        self.add_widget(b1)


    def change(self,*args):
        sm.doc_file_manager()

    def changet(self,*args):
        sm.file_manager_open()

class ContentNavigationDrawer(BoxLayout,Thread):
    screen_manager = ObjectProperty()
    nav_drawer = ObjectProperty()


class Ink(MDApp,Thread):

    

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.theme_cls.theme_style = "Dark"
        
        

    def events(self, instance, keyboard, keycode, text, modifiers):

        if keyboard in (1001, 27):
            if self.manager_open:
                self.file_manager.back()
        return True

    def change_screen(self,*args):
        sm.current='Main'

    def build(self):

        self.Manager=ScreenManager()
        self.theme_cls.primary_palette = "BlueGray"
        return Builder.load_string(
            '''
<MainScreen>:
    name: 'Main'
    MDNavigationLayout:
        screen_manager:screen_manager
        Manager:
            id : screen_manager
            MenuScreen:
                name: 'Menu'
                BoxLayout:
                    
                    orientation:'vertical'
            
                    MDToolbar:
                        
                        id: toolbar

                        title: 'INK'
                        md_bg_color: 0.125, 0.125, 0.125, 1
                        specific_text_color: 1, 1, 1, 1
                        left_action_items: [["menu", lambda x: nav_drawer.set_state('toggle')]]
                        
                    MDBottomNavigation:
                    
                        panel_color: 0.125, 0.125, 0.125, 1
                        text_color_active: 0.5,0.75,1, 1
                        

                        MDBottomNavigationItem:
                            name: 'screen 1'
                            text: 'Home'
                            icon: 'home'

                            Front_BoxLayout:    
                

                        MDBottomNavigationItem:
                            name: 'screen 3'
                            text: 'Font'
                            icon: 'pen'

                            
                            Writing:
                                    
                            
                            MDFloatingActionButton:
                                icon: "plus"
                                pos_hint:{'right' : 0.95,'top':0.15}
                                on_release: root.py_test()
                                opposite_colors: True

                        MDBottomNavigationItem:
                            name: 'screen 2'
                            text:'setting'
                            icon: 'brightness-5'
                            
                            ScrollView:
                                Setting:

                            
                
            About_Screen:
                name: "About"
                BoxLayout:
                    orientation:'vertical'
                
                    MDToolbar:
                        id: toolbar
                        title: 'Developer'
                        md_bg_color: 0.125, 0.125, 0.125, 1
                        specific_text_color: 1, 1, 1, 1
                        left_action_items: [["menu", lambda x: nav_drawer.set_state('toggle')]]
                        right_action_items: [["arrow-left-thick", lambda x: screen_manager.change_screen()]]

                    Label:
                        text: 'Developed by Satyam Lohiya'
                        

        MDNavigationDrawer:
            type: "modal"
            id: nav_drawer

            ContentNavigationDrawer:
                screen_manager: screen_manager
                nav_drawer: nav_drawer


<toolbar>:
    id: toolbar
    title: 'IMAGE'
    md_bg_color: 0.125, 0.125, 0.125, 1
    specific_text_color: 1, 1, 1, 1
    left_action_items: [["arrow-left-thick", lambda x: app.change_screen()]]


<Btoolbar>:
    id: Btoolbar
    title: 'BG-Image'
    md_bg_color: 0.125, 0.125, 0.125, 1
    specific_text_color: 1, 1, 1, 1
    left_action_items: [["arrow-left-thick", lambda x: app.change_screen()]]



<ContentNavigationDrawer>:
    padding: "8dp"
    spacing: "8dp"

    ScrollView:

        MDList:

            OneLineIconListItem:
                text: "Home"
                on_release:
                    root.nav_drawer.set_state("close")
                    root.screen_manager.current = "Menu"
                IconLeftWidget:
                    icon: 'home'

            OneLineIconListItem:
                text: "Developer"
                on_release:
                    root.nav_drawer.set_state("close")
                    root.screen_manager.current = "About"
                IconLeftWidget:
                    icon: 'account' 

<BGimage>:
    name: "BGimage"
    
    MDFloatingActionButton:
        icon: "plus"
        pos_hint:{'right' : 0.95,'top':0.15}
        on_release: root.file_manager_open()
        opposite_colors: True

<TableImage>:
    name: "TableImage"
    MDFloatingActionButton:
        icon: "plus"
        pos_hint:{'right' : 0.95,'top':0.15}
        on_release: root.file_manager_open()
        opposite_colors: True

Scr:
    id: 'scr'
    MainScreen:

    BGimage:
        name: "BGimage"
    
    TableImage:
        name: "TableImage"
        
    TableImagePreview:
        name: "TablePreview"
    
    Image_Choose:
        name: "Setting"
        

'''
        )

Ink().run()